"""Reusable Word-document builder for the Demo 06 cybersecurity deliverables.

Wraps python-docx with a small, opinionated styling layer so each
deliverable (architecture article, deployment guide, test report, pitch
deck) shares one consistent visual language across English and Ukrainian
editions. PDF rendering is handled separately by LibreOffice headless.
"""

from __future__ import annotations

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Emu


# Brand palette -----------------------------------------------------------
INK = RGBColor(0x14, 0x1B, 0x2E)        # near-black navy for body text
NAVY = RGBColor(0x10, 0x2A, 0x54)       # deep navy for H1
BLUE = RGBColor(0x1C, 0x5D, 0x99)       # mid blue for H2
TEAL = RGBColor(0x0E, 0x7C, 0x86)       # teal accent for H3
GREY = RGBColor(0x55, 0x5F, 0x70)       # muted grey for captions
LIGHT = RGBColor(0x8A, 0x93, 0xA6)      # light grey
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GOLD = RGBColor(0xC8, 0x8A, 0x1A)       # accent gold
DANGER = RGBColor(0xB1, 0x2A, 0x2A)     # red for warnings
GREEN = RGBColor(0x1F, 0x7A, 0x3D)      # green for success

BODY_FONT = "Calibri"
HEAD_FONT = "Calibri"
MONO_FONT = "Consolas"


def _shade(cell, hex_color: str) -> None:
    """Apply a solid background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=110, right=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for name, value in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tc_pr.append(margins)


def _no_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "none")
        borders.append(node)
    tbl_pr.append(borders)


def _hairline_borders(table, color="D6DBE5") -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)
    tbl_pr.append(borders)


def _vertical_center(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    valign = OxmlElement("w:vAlign")
    valign.set(qn("w:val"), "center")
    tc_pr.append(valign)


class DocBuilder:
    """Thin convenience layer over python-docx with house styles."""

    def __init__(self, accent: RGBColor = BLUE):
        self.doc = Document()
        self.accent = accent
        self._configure_base_styles()
        self._configure_page()

    # -- base configuration ------------------------------------------------
    def _configure_base_styles(self) -> None:
        normal = self.doc.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = INK
        pf = normal.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.18

    def _configure_page(self, landscape: bool = False) -> None:
        section = self.doc.sections[0]
        if landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
        else:
            section.page_width = Inches(8.27)   # A4
            section.page_height = Inches(11.69)
            section.left_margin = Inches(0.95)
            section.right_margin = Inches(0.95)
            section.top_margin = Inches(0.9)
            section.bottom_margin = Inches(0.9)

    # -- footer ------------------------------------------------------------
    def add_footer(self, left_text: str) -> None:
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(left_text)
        run.font.size = Pt(8)
        run.font.color.rgb = LIGHT
        run.font.name = BODY_FONT
        # page number field
        p.add_run("    ").font.size = Pt(8)
        self._add_page_field(p)

    def _add_page_field(self, paragraph) -> None:
        run = paragraph.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.size = Pt(8)
        run.font.color.rgb = LIGHT

    # -- primitives --------------------------------------------------------
    def spacer(self, pts: float = 6) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run("")
        run.font.size = Pt(pts)

    def h1(self, text: str, number: str | None = None) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        if number:
            r = p.add_run(f"{number}  ")
            r.font.size = Pt(17)
            r.font.bold = True
            r.font.color.rgb = self.accent
            r.font.name = HEAD_FONT
        run = p.add_run(text)
        run.font.size = Pt(17)
        run.font.bold = True
        run.font.color.rgb = NAVY
        run.font.name = HEAD_FONT
        self._bottom_rule(p, self.accent)

    def _bottom_rule(self, paragraph, color: RGBColor) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "10")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), f"{color}")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def h2(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = BLUE
        run.font.name = HEAD_FONT

    def h3(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = TEAL
        run.font.name = HEAD_FONT

    def para(self, text: str, size: float = 10.5, italic: bool = False,
             color: RGBColor = INK, align=None, space_after: float = 6) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        if align is not None:
            p.alignment = align
        self._emit_runs(p, text, size=size, italic=italic, color=color)

    def _emit_runs(self, paragraph, text: str, size: float, italic: bool, color: RGBColor) -> None:
        """Render text supporting **bold** and `code` inline markup."""
        import re
        tokens = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
        for token in tokens:
            if not token:
                continue
            if token.startswith("**") and token.endswith("**"):
                run = paragraph.add_run(token[2:-2])
                run.font.bold = True
            elif token.startswith("`") and token.endswith("`"):
                run = paragraph.add_run(token[1:-1])
                run.font.name = MONO_FONT
                run.font.size = Pt(size - 0.5)
                run.font.color.rgb = NAVY
                continue
            else:
                run = paragraph.add_run(token)
            run.font.size = Pt(size)
            run.font.italic = italic
            run.font.color.rgb = color
            run.font.name = BODY_FONT

    def bullet(self, text: str, level: int = 0, size: float = 10.5) -> None:
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.28 + 0.25 * level)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        self._emit_runs(p, text, size=size, italic=False, color=INK)

    def numbered(self, text: str, index: int = 1, size: float = 10.5) -> None:
        # Manual numbering: Word's "List Number" style continues globally and
        # will not restart per section, so we render the number ourselves.
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        marker = p.add_run(f"{index}.  ")
        marker.font.size = Pt(size)
        marker.font.bold = True
        marker.font.color.rgb = self.accent
        marker.font.name = HEAD_FONT
        self._emit_runs(p, text, size=size, italic=False, color=INK)

    def code(self, text: str, size: float = 9) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _no_table_borders(table)
        cell = table.cell(0, 0)
        _shade(cell, "0E1A2B")
        _set_cell_margins(cell, top=90, bottom=90, left=140, right=140)
        first = True
        for line in text.split("\n"):
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.06
            run = p.add_run(line if line else " ")
            run.font.name = MONO_FONT
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(0xD6, 0xE7, 0xC8)
        self.spacer(4)

    def callout(self, title: str, body: str, kind: str = "info") -> None:
        colors = {
            "info": ("EAF1F9", "1C5D99"),
            "success": ("E8F3EC", "1F7A3D"),
            "warning": ("FBF1E0", "C88A1A"),
            "danger": ("F8E9E9", "B12A2A"),
        }
        fill, bar = colors.get(kind, colors["info"])
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _hairline_borders(table, color=fill)
        cell = table.cell(0, 0)
        _shade(cell, fill)
        _set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
        self._left_accent(cell, bar)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(title)
        r.font.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(int(bar[0:2], 16), int(bar[2:4], 16), int(bar[4:6], 16))
        r.font.name = HEAD_FONT
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        self._emit_runs(p2, body, size=10, italic=False, color=INK)
        self.spacer(6)

    def _left_accent(self, cell, hex_color: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")
        left.set(qn("w:space"), "0")
        left.set(qn("w:color"), hex_color)
        borders.append(left)
        tc_pr.append(borders)

    def table(self, headers: list[str], rows: list[list[str]],
              widths: list[float] | None = None, font_size: float = 9.5,
              header_fill: str | None = None) -> None:
        header_fill = header_fill or f"{self.accent}"
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _hairline_borders(t)
        hdr = t.rows[0].cells
        for i, head in enumerate(headers):
            _shade(hdr[i], header_fill)
            _set_cell_margins(hdr[i])
            _vertical_center(hdr[i])
            p = hdr[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(head)
            r.font.bold = True
            r.font.size = Pt(font_size)
            r.font.color.rgb = WHITE
            r.font.name = HEAD_FONT
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            stripe = "F4F7FB" if ri % 2 == 0 else "FFFFFF"
            for ci, value in enumerate(row):
                _shade(cells[ci], stripe)
                _set_cell_margins(cells[ci])
                _vertical_center(cells[ci])
                p = cells[ci].paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                self._emit_runs(p, value, size=font_size, italic=False, color=INK)
        if widths:
            for ci, w in enumerate(widths):
                for row in t.rows:
                    row.cells[ci].width = Inches(w)
        self.spacer(6)

    def page_break(self) -> None:
        self.doc.add_page_break()

    # -- cover page --------------------------------------------------------
    def cover(self, kicker: str, title: str, subtitle: str,
              meta_rows: list[tuple[str, str]], doc_type: str,
              tagline: str | None = None) -> None:
        # top color band
        band = self.doc.add_table(rows=1, cols=1)
        _no_table_borders(band)
        bcell = band.cell(0, 0)
        _shade(bcell, "102A54")
        _set_cell_margins(bcell, top=260, bottom=260, left=240, right=240)
        p = bcell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(kicker.upper())
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x8F, 0xC2, 0xF0)
        r.font.name = HEAD_FONT
        p2 = bcell.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(doc_type.upper())
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0xB9, 0xC6, 0xDA)
        r2.font.name = HEAD_FONT

        self.spacer(20)
        tp = self.doc.add_paragraph()
        tp.paragraph_format.space_after = Pt(4)
        tr = tp.add_run(title)
        tr.font.size = Pt(30)
        tr.font.bold = True
        tr.font.color.rgb = NAVY
        tr.font.name = HEAD_FONT

        sp = self.doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(10)
        sr = sp.add_run(subtitle)
        sr.font.size = Pt(14)
        sr.font.color.rgb = BLUE
        sr.font.name = HEAD_FONT

        # accent rule
        rule = self.doc.add_paragraph()
        self._bottom_rule(rule, self.accent)
        self.spacer(8)

        if tagline:
            self.para(tagline, size=11.5, italic=True, color=GREY)
            self.spacer(10)

        # meta block
        t = self.doc.add_table(rows=0, cols=2)
        _no_table_borders(t)
        for label, value in meta_rows:
            cells = t.add_row().cells
            _set_cell_margins(cells[0], top=40, bottom=40, left=0, right=120)
            _set_cell_margins(cells[1], top=40, bottom=40, left=0, right=0)
            pl = cells[0].paragraphs[0]
            pl.paragraph_format.space_after = Pt(0)
            rl = pl.add_run(label)
            rl.font.size = Pt(9.5)
            rl.font.bold = True
            rl.font.color.rgb = GREY
            rl.font.name = HEAD_FONT
            pv = cells[1].paragraphs[0]
            pv.paragraph_format.space_after = Pt(0)
            rv = pv.add_run(value)
            rv.font.size = Pt(9.5)
            rv.font.color.rgb = INK
            rv.font.name = BODY_FONT
        cells[0].width = Inches(1.7)
        cells[1].width = Inches(4.6)
        self.page_break()

    def toc(self, title: str, entries: list[str]) -> None:
        self.h1(title)
        for i, entry in enumerate(entries, start=1):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.1)
            r = p.add_run(f"{i:>2}.  ")
            r.font.size = Pt(10.5)
            r.font.bold = True
            r.font.color.rgb = self.accent
            r.font.name = HEAD_FONT
            r2 = p.add_run(entry)
            r2.font.size = Pt(10.5)
            r2.font.color.rgb = INK
            r2.font.name = BODY_FONT
        self.page_break()

    def save(self, path: str) -> None:
        self.doc.save(path)


# -- Pitch-deck (slide-style, landscape) ----------------------------------
class SlideBuilder:
    """Landscape, one-slide-per-page deck rendered as a DOCX/PDF."""

    def __init__(self):
        self.doc = Document()
        normal = self.doc.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(14)
        normal.font.color.rgb = INK
        section = self.doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(13.333)
        section.page_height = Inches(7.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.5)
        self._first = True

    def _new_slide(self) -> None:
        if not self._first:
            self.doc.add_page_break()
        self._first = False

    def cover_slide(self, kicker: str, title: str, subtitle: str, footer: str) -> None:
        self._new_slide()
        band = self.doc.add_table(rows=1, cols=1)
        _no_table_borders(band)
        cell = band.cell(0, 0)
        _shade(cell, "102A54")
        _set_cell_margins(cell, top=900, bottom=900, left=420, right=420)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(kicker.upper())
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x8F, 0xC2, 0xF0)
        tp = cell.add_paragraph()
        tp.paragraph_format.space_after = Pt(8)
        tr = tp.add_run(title)
        tr.font.size = Pt(40)
        tr.font.bold = True
        tr.font.color.rgb = WHITE
        sp = cell.add_paragraph()
        sr = sp.add_run(subtitle)
        sr.font.size = Pt(18)
        sr.font.color.rgb = RGBColor(0xCF, 0xDC, 0xEE)
        fp = self.doc.add_paragraph()
        fp.paragraph_format.space_before = Pt(16)
        fr = fp.add_run(footer)
        fr.font.size = Pt(11)
        fr.font.color.rgb = GREY

    def _slide_title(self, index: str, title: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{index}  ")
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = TEAL
        r2 = p.add_run(title)
        r2.font.size = Pt(26)
        r2.font.bold = True
        r2.font.color.rgb = NAVY
        rule = self.doc.add_paragraph()
        p_pr = rule._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "1C5D99")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
        rule.paragraph_format.space_after = Pt(8)

    def bullet_slide(self, index: str, title: str, lead: str | None,
                     bullets: list[str], kind: str = "plain") -> None:
        self._new_slide()
        self._slide_title(index, title)
        if lead:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(lead)
            r.font.size = Pt(15)
            r.font.italic = True
            r.font.color.rgb = BLUE
        for b in bullets:
            self._deck_bullet(b)

    def _deck_bullet(self, text: str) -> None:
        import re
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.line_spacing = 1.1
        dot = p.add_run("▸  ")
        dot.font.size = Pt(14)
        dot.font.color.rgb = TEAL
        dot.font.bold = True
        tokens = re.split(r"(\*\*.*?\*\*)", text)
        for token in tokens:
            if not token:
                continue
            if token.startswith("**") and token.endswith("**"):
                run = p.add_run(token[2:-2])
                run.font.bold = True
                run.font.color.rgb = NAVY
            else:
                run = p.add_run(token)
                run.font.color.rgb = INK
            run.font.size = Pt(14.5)

    def metric_slide(self, index: str, title: str, lead: str,
                     metrics: list[tuple[str, str]]) -> None:
        self._new_slide()
        self._slide_title(index, title)
        if lead:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            r = p.add_run(lead)
            r.font.size = Pt(15)
            r.font.italic = True
            r.font.color.rgb = BLUE
        cols = len(metrics)
        t = self.doc.add_table(rows=2, cols=cols)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _no_table_borders(t)
        for i, (big, small) in enumerate(metrics):
            top = t.cell(0, i)
            _shade(top, "EAF1F9")
            _set_cell_margins(top, top=180, bottom=60, left=120, right=120)
            p = top.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(big)
            r.font.size = Pt(30)
            r.font.bold = True
            r.font.color.rgb = NAVY
            bot = t.cell(1, i)
            _shade(bot, "EAF1F9")
            _set_cell_margins(bot, top=0, bottom=180, left=120, right=120)
            p2 = bot.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(small)
            r2.font.size = Pt(11)
            r2.font.color.rgb = GREY

    def two_col_slide(self, index: str, title: str,
                      left_head: str, left_items: list[str],
                      right_head: str, right_items: list[str]) -> None:
        self._new_slide()
        self._slide_title(index, title)
        t = self.doc.add_table(rows=1, cols=2)
        _no_table_borders(t)
        left, right = t.cell(0, 0), t.cell(0, 1)
        _shade(left, "F8E9E9")
        _shade(right, "E8F3EC")
        _set_cell_margins(left, top=140, bottom=160, left=180, right=180)
        _set_cell_margins(right, top=140, bottom=160, left=180, right=180)
        for cell, head, items, color in (
            (left, left_head, left_items, DANGER),
            (right, right_head, right_items, GREEN),
        ):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(head)
            r.font.bold = True
            r.font.size = Pt(15)
            r.font.color.rgb = color
            for item in items:
                ip = cell.add_paragraph()
                ip.paragraph_format.space_after = Pt(5)
                ip.paragraph_format.left_indent = Inches(0.1)
                d = ip.add_run("• ")
                d.font.color.rgb = color
                d.font.size = Pt(13)
                tr = ip.add_run(item)
                tr.font.size = Pt(12.5)
                tr.font.color.rgb = INK
        left.width = Inches(5.9)
        right.width = Inches(5.9)

    def closing_slide(self, title: str, lines: list[str], cta: str, contact: list[str]) -> None:
        self._new_slide()
        band = self.doc.add_table(rows=1, cols=1)
        _no_table_borders(band)
        cell = band.cell(0, 0)
        _shade(cell, "102A54")
        _set_cell_margins(cell, top=520, bottom=520, left=420, right=420)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(title)
        r.font.size = Pt(32)
        r.font.bold = True
        r.font.color.rgb = WHITE
        for line in lines:
            lp = cell.add_paragraph()
            lp.paragraph_format.space_after = Pt(4)
            lr = lp.add_run(line)
            lr.font.size = Pt(15)
            lr.font.color.rgb = RGBColor(0xCF, 0xDC, 0xEE)
        cp = cell.add_paragraph()
        cp.paragraph_format.space_before = Pt(10)
        cr = cp.add_run(cta)
        cr.font.size = Pt(18)
        cr.font.bold = True
        cr.font.color.rgb = RGBColor(0x8F, 0xC2, 0xF0)
        cp2 = self.doc.add_paragraph()
        cp2.paragraph_format.space_before = Pt(14)
        for i, c in enumerate(contact):
            run = cp2.add_run(("" if i == 0 else "    |    ") + c)
            run.font.size = Pt(11)
            run.font.color.rgb = GREY

    def save(self, path: str) -> None:
        self.doc.save(path)
